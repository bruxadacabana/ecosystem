"""
Carregadores de documentos para diferentes formatos.
Suporta: PDF, DOCX, TXT, MD, EPUB, VTT, SRT.
Para .md em vault Obsidian: extrai frontmatter YAML e wikilinks.
Para .vtt/.srt e .md/.txt do Hermes: detectado como transcrição automaticamente.
"""
from __future__ import annotations

import os
import re

from langchain_core.documents import Document
from langchain_community.document_loaders import (
    PyPDFLoader,
    Docx2txtLoader,
    TextLoader,
)

from .errors import DocumentLoadError, FrontmatterParseError, UnsupportedFormatError


_SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt", ".md", ".epub",
                         ".mobi", ".azw", ".azw3",
                         ".jpg", ".jpeg", ".png", ".webp",
                         ".vtt", ".srt"}

_IMAGE_EXTENSIONS      = frozenset({".jpg", ".jpeg", ".png", ".webp"})
_TRANSCRIPT_FILE_EXTS  = frozenset({".vtt", ".srt"})

# Cobre os 4 formatos de wikilink do Obsidian:
# [[nota]], [[nota|alias]], [[nota#secção]], [[nota#secção|alias]]
_WIKILINK_RE = re.compile(r"\[\[([^\]|#]+)(?:#[^\]|]+)?(?:\|([^\]]+))?\]\]")

# Diretórios a ignorar ao percorrer um vault Obsidian
_OBSIDIAN_IGNORE = {".obsidian", "templates", "attachments", ".trash"}

# Detecta linha de timing VTT (00:00:00.000 --> ...) e SRT (00:00:00,000 --> ...)
_VTT_TIMING_RE  = re.compile(r"^\d{1,2}:\d{2}[:.]\d{2}[,.]\d{3}\s+-->")
_VTT_CUE_NUM_RE = re.compile(r"^\d+$")
# Detecta timestamp inline Hermes: **[MM:SS]** texto
_HERMES_STAMP_RE = re.compile(r"^\*\*\[\d{2}:\d{2}\]\*\*")


def is_transcript_file(file_path: str) -> bool:
    """Detecta se um arquivo é uma transcrição (Hermes .md/.txt ou legendas .vtt/.srt).

    Critérios para .md/.txt:
    - Frontmatter YAML com chave ``duration:`` (padrão Hermes)
    - Heading ``## Transcrição``
    - 3 ou mais linhas com timestamps ``**[MM:SS]**``
    """
    ext = os.path.splitext(file_path.lower())[1]
    if ext in _TRANSCRIPT_FILE_EXTS:
        return True
    if ext not in (".md", ".txt"):
        return False
    try:
        with open(file_path, encoding="utf-8", errors="ignore") as fh:
            lines = [fh.readline() for _ in range(50)]
    except OSError:
        return False
    in_fm = bool(lines) and lines[0].strip() == "---"
    fm_closed = False
    stamp_count = 0
    for i, line in enumerate(lines):
        s = line.strip()
        if i == 0 and s == "---":
            continue
        if in_fm and not fm_closed:
            if s == "---":
                fm_closed = True
            elif s.startswith("duration:"):
                return True
            continue
        if s == "## Transcrição":
            return True
        if _HERMES_STAMP_RE.match(s):
            stamp_count += 1
            if stamp_count >= 3:
                return True
    return False


def _load_vtt(file_path: str) -> list[Document]:
    """Carrega .vtt ou .srt extraindo apenas o texto dos cues, sem timestamps."""
    try:
        raw = open(file_path, encoding="utf-8", errors="ignore").read()
    except OSError as exc:
        raise DocumentLoadError(file_path, str(exc)) from exc
    text_lines: list[str] = []
    for line in raw.splitlines():
        s = line.strip()
        if not s or s == "WEBVTT" or _VTT_TIMING_RE.match(s) or _VTT_CUE_NUM_RE.match(s):
            continue
        cleaned = re.sub(r"<[^>]+>", "", s)
        if cleaned:
            text_lines.append(cleaned)
    title = os.path.splitext(os.path.basename(file_path))[0]
    return [Document(
        page_content="\n".join(text_lines),
        metadata={"source": file_path, "title": title},
    )]


def load_documents(
    directory: str,
    source_type: str = "library",
    ocr_model: str = "",
) -> tuple[list[Document], list[DocumentLoadError]]:
    """
    Carrega todos os documentos suportados de um diretório (recursivo).
    source_type "vault" ativa o loader Obsidian para arquivos .md.
    ocr_model: modelo Ollama vision para fallback OCR de imagens (vazio = Tesseract only).

    Retorna (documentos_carregados, erros_por_arquivo).
    """
    if not os.path.exists(directory):
        raise FileNotFoundError(f"Diretório não encontrado: {directory}")

    documents: list[Document] = []
    errors: list[DocumentLoadError] = []

    ignore_dirs = {".mnemosyne"}
    if source_type == "vault":
        ignore_dirs |= _OBSIDIAN_IGNORE

    for root, dirs, files in os.walk(directory):
        dirs[:] = [d for d in dirs if d not in ignore_dirs]
        for filename in sorted(files):
            _, ext = os.path.splitext(filename.lower())
            if ext not in _SUPPORTED_EXTENSIONS:
                continue
            file_path = os.path.join(root, filename)
            try:
                docs = _load_file(file_path, source_type=source_type, ocr_model=ocr_model)
                documents.extend(docs)
            except DocumentLoadError as exc:
                errors.append(exc)

    return documents, errors


def load_single_file(
    file_path: str,
    source_type: str = "library",
    ocr_model: str = "",
) -> list[Document]:
    """
    Carrega um único arquivo.
    ocr_model: modelo Ollama vision para fallback OCR de imagens (vazio = Tesseract only).

    Raises:
        FileNotFoundError: se o arquivo não existir.
        UnsupportedFormatError: se o formato não for suportado.
        DocumentLoadError: se a leitura falhar.
    """
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"Arquivo não encontrado: {file_path}")
    return _load_file(file_path, source_type=source_type, ocr_model=ocr_model)


def _load_file(
    file_path: str,
    source_type: str = "library",
    ocr_model: str = "",
) -> list[Document]:
    """
    Despacha para o loader correto e define source_type no metadata.

    ocr_model: modelo Ollama vision para fallback OCR de imagens (vazio = Tesseract only).

    Raises:
        UnsupportedFormatError: se a extensão não for suportada.
        DocumentLoadError: em qualquer falha de leitura.
    """
    _, ext = os.path.splitext(file_path.lower())

    try:
        if ext == ".pdf":
            docs = _load_pdf(file_path)
        elif ext == ".docx":
            docs = _load_docx(file_path)
        elif ext in _TRANSCRIPT_FILE_EXTS:
            docs = _load_vtt(file_path)
            source_type = "transcript"
        elif ext in (".txt", ".md"):
            if source_type == "vault":
                docs = _load_obsidian_note(file_path)
            else:
                if is_transcript_file(file_path):
                    source_type = "transcript"
                    docs = TextLoader(file_path, encoding="utf-8").load()
                elif ext == ".md":
                    docs = _load_library_md(file_path)
                    # Outputs do Studio têm frontmatter `source: mnemosyne_studio`
                    if docs and docs[0].metadata.get("fm_source") == "mnemosyne_studio":
                        source_type = "thought"
                else:
                    docs = TextLoader(file_path, encoding="utf-8").load()
        elif ext == ".epub":
            docs = _load_epub(file_path)
            source_type = "book"
        elif ext in (".mobi", ".azw", ".azw3"):
            docs = _load_mobi(file_path)
        elif ext in _IMAGE_EXTENSIONS:
            docs = _load_image(file_path, ocr_model=ocr_model)
        else:
            raise UnsupportedFormatError(file_path)
    except (DocumentLoadError, UnsupportedFormatError):
        raise
    except OSError as exc:
        raise DocumentLoadError(file_path, str(exc)) from exc
    except Exception as exc:
        raise DocumentLoadError(file_path, str(exc)) from exc

    for doc in docs:
        doc.metadata["source_type"] = source_type
    return docs


def _load_pdf(file_path: str) -> list[Document]:
    """Carrega PDF e garante metadata title/author nas páginas."""
    docs = PyPDFLoader(file_path).load()
    # PyPDFLoader já inclui metadata do PDF; apenas normalizar chaves
    for doc in docs:
        if "title" not in doc.metadata:
            doc.metadata["title"] = os.path.splitext(os.path.basename(file_path))[0]
        if "author" not in doc.metadata:
            doc.metadata["author"] = ""
    return docs


def _load_docx(file_path: str) -> list[Document]:
    """Carrega DOCX e extrai title/author das propriedades do documento."""
    docs = Docx2txtLoader(file_path).load()
    title = os.path.splitext(os.path.basename(file_path))[0]
    author = ""
    try:
        import docx as _docx
        d = _docx.Document(file_path)
        props = d.core_properties
        if props.title:
            title = props.title
        if props.author:
            author = props.author
    except Exception:
        pass
    for doc in docs:
        doc.metadata.setdefault("title", title)
        doc.metadata.setdefault("author", author)
    return docs


def extract_wikilinks(text: str) -> list[str]:
    """Extrai nomes de notas de todos os wikilinks presentes no texto."""
    return [m.group(1).strip() for m in _WIKILINK_RE.finditer(text)]


def _load_obsidian_note(file_path: str) -> list[Document]:
    """
    Carrega uma nota Obsidian (.md) com:
    - Frontmatter YAML (title, tags, aliases)
    - Wikilinks extraídos como metadata
    - Chunking por cabeçalho ##: uma nota vira 1 ou N documentos por secção
    - Notas com menos de 50 chars de corpo são ignoradas

    Raises:
        DocumentLoadError: se o arquivo não puder ser lido.
        FrontmatterParseError: se o frontmatter for inválido (capturado internamente).
    """
    try:
        raw_text = open(file_path, encoding="utf-8", errors="ignore").read()
    except OSError as exc:
        raise DocumentLoadError(file_path, str(exc)) from exc

    # Extrair frontmatter
    fm: dict = {}
    body = raw_text
    try:
        import frontmatter as _fm  # python-frontmatter
        post = _fm.loads(raw_text)
        fm = dict(post.metadata)
        body = post.content
    except Exception:
        pass  # sem frontmatter ou erro de parse: usa texto bruto

    title = str(fm.get("title", "")) or os.path.splitext(os.path.basename(file_path))[0]
    tags: list[str] = []
    raw_tags = fm.get("tags", [])
    if isinstance(raw_tags, list):
        tags = [str(t) for t in raw_tags]
    elif isinstance(raw_tags, str):
        tags = [t.strip() for t in raw_tags.split(",") if t.strip()]

    aliases: list[str] = []
    raw_aliases = fm.get("aliases", [])
    if isinstance(raw_aliases, list):
        aliases = [str(a) for a in raw_aliases]
    elif isinstance(raw_aliases, str):
        aliases = [raw_aliases]

    wikilinks = extract_wikilinks(body)

    base_meta = {
        "source": file_path,
        "title": title,
        "tags": ", ".join(tags),
        "aliases": ", ".join(aliases),
        "wikilinks": ", ".join(wikilinks),
    }

    # Chunking por cabeçalho ## (nível 2)
    sections = _split_by_heading(body, title)

    documents: list[Document] = []
    for sec_title, sec_body in sections:
        sec_body = sec_body.strip()
        if len(sec_body) < 50:
            continue
        meta = {**base_meta, "section": sec_title}
        documents.append(Document(page_content=sec_body, metadata=meta))

    return documents


def _split_by_heading(text: str, fallback_title: str) -> list[tuple[str, str]]:
    """
    Divide texto por cabeçalhos ## (nível 2).
    Retorna lista de (título_da_secção, corpo).
    Se não houver cabeçalhos, retorna o texto inteiro com fallback_title.
    """
    lines = text.splitlines()
    sections: list[tuple[str, str]] = []
    current_title = fallback_title
    current_lines: list[str] = []

    for line in lines:
        if line.startswith("## "):
            if current_lines:
                sections.append((current_title, "\n".join(current_lines)))
            current_title = line[3:].strip()
            current_lines = []
        else:
            current_lines.append(line)

    if current_lines:
        sections.append((current_title, "\n".join(current_lines)))

    return sections if sections else [(fallback_title, text)]


def _load_image(file_path: str, ocr_model: str = "") -> list[Document]:
    """
    Extrai texto de uma imagem (JPG, PNG, WebP) para indexação RAG.

    Estratégia em duas camadas:
      1. Tesseract via pytesseract (rápido, sem GPU, <100 MB RAM) — caminho principal.
      2. Fallback para Ollama vision via /api/generate com imagem em base64,
         usado quando Tesseract falha ou quando ocr_model está configurado.

    O texto extraído vira um Document com metadata ocr_engine informando qual
    camada foi usada ("tesseract" ou "ollama:{model}").

    Raises:
        DocumentLoadError: se ambas as camadas falharem ou nenhuma estiver disponível.
    """
    import base64
    from pathlib import Path as _Path

    title = _Path(file_path).stem
    text = ""
    engine = ""

    # Camada 1: Tesseract
    try:
        import pytesseract  # type: ignore[import]
        from PIL import Image  # type: ignore[import]

        img = Image.open(file_path)
        # Converter para RGB se necessário (PNG com canal alfa não aceite pelo Tesseract)
        if img.mode not in ("RGB", "L"):
            img = img.convert("RGB")
        text = pytesseract.image_to_string(img, lang="por+eng")
        engine = "tesseract"
    except ImportError:
        pass  # pytesseract/Pillow não instalado — tentar Ollama
    except Exception:
        pass  # Tesseract falhou na imagem específica — tentar Ollama

    # Camada 2: Ollama vision — usado se Tesseract falhou OU se ocr_model configurado
    if (not text.strip() or ocr_model) and ocr_model:
        try:
            import httpx
            from .ollama_client import _BASE_URL as _OLLAMA_BASE

            with open(file_path, "rb") as fh:
                img_b64 = base64.b64encode(fh.read()).decode()

            resp = httpx.post(
                f"{_OLLAMA_BASE}/api/generate",
                json={
                    "model": ocr_model,
                    "prompt": (
                        "Extraia TODO o texto visível nesta imagem. "
                        "Se houver tabelas, preserva a estrutura em Markdown. "
                        "Se não houver texto, descreve o conteúdo da imagem em português."
                    ),
                    "images": [img_b64],
                    "stream": False,
                    "temperature": 0,
                    "keep_alive": "10m",
                },
                timeout=120.0,
            )
            resp.raise_for_status()
            ollama_text = resp.json().get("response", "").strip()
            if ollama_text:
                text = ollama_text
                engine = f"ollama:{ocr_model}"
        except Exception:
            pass

    if not text.strip():
        raise DocumentLoadError(
            file_path,
            "Não foi possível extrair texto da imagem. "
            "Verifique se pytesseract está instalado ou configure image_ocr_model.",
        )

    return [Document(
        page_content=text.strip(),
        metadata={"source": file_path, "title": title, "ocr_engine": engine},
    )]


def _load_mobi(file_path: str) -> list[Document]:
    """
    Carrega MOBI, AZW ou AZW3 via biblioteca `mobi` (KindleUnpack).

    A saída da extração pode ser:
      - index.html (MOBI puro) → BeautifulSoup
      - *.epub (AZW3/KF8)     → reutiliza _load_epub()
      - *.pdf (Print Replica) → PyPDFLoader

    Arquivos com DRM ativo geram output corrompido; deteta-se pelo tamanho
    do texto extraído e retorna DocumentLoadError informativo.

    Raises:
        DocumentLoadError: se `mobi` não estiver instalado, DRM detectado, ou falha de leitura.
    """
    import tempfile
    from pathlib import Path as _Path

    try:
        import mobi as _mobi  # type: ignore[import]
    except ImportError as exc:
        raise DocumentLoadError(
            file_path,
            "Dependência em falta para MOBI/AZW: mobi. Instale com: pip install mobi",
        ) from exc

    title = _Path(file_path).stem
    author = ""

    with tempfile.TemporaryDirectory() as tmpdir:
        try:
            _, result_path = _mobi.extract(file_path, tmpdir)
        except Exception as exc:
            raise DocumentLoadError(file_path, f"Falha ao extrair MOBI/AZW: {exc}") from exc

        result = _Path(result_path)

        # AZW3/KF8 → EPUB gerado dentro do tmpdir
        epub_files = list(result.parent.glob("**/*.epub"))
        if epub_files:
            try:
                return _load_epub(str(epub_files[0]))
            except DocumentLoadError as exc:
                raise DocumentLoadError(file_path, str(exc)) from exc

        # Print Replica AZW → PDF
        pdf_files = list(result.parent.glob("**/*.pdf"))
        if pdf_files:
            return _load_pdf(str(pdf_files[0]))

        # MOBI puro → index.html
        html_candidates = [result] if result.suffix.lower() in (".html", ".htm") else []
        if not html_candidates:
            html_candidates = list(result.parent.glob("**/index.html"))
        if not html_candidates:
            html_candidates = list(result.parent.glob("**/*.html"))
        if not html_candidates:
            raise DocumentLoadError(
                file_path,
                "Nenhum conteúdo extraído — arquivo provavelmente protegido por DRM.",
            )

        html_path = html_candidates[0]
        raw_html = html_path.read_text(encoding="utf-8", errors="ignore")

        try:
            from bs4 import BeautifulSoup
        except ImportError as exc:
            raise DocumentLoadError(
                file_path, "beautifulsoup4 não instalado: pip install beautifulsoup4"
            ) from exc

        soup = BeautifulSoup(raw_html, "lxml")
        text = soup.get_text(separator="\n", strip=True)

        if len(text) < 50:
            raise DocumentLoadError(
                file_path,
                "Texto extraído muito curto — arquivo provavelmente protegido por DRM.",
            )

        # Tentar extrair título/autor das meta tags
        meta_title = soup.find("meta", attrs={"name": "title"})
        if meta_title:
            title = meta_title.get("content", title) or title
        title_tag = soup.find("title")
        if title_tag and title_tag.string:
            title = title_tag.string.strip() or title
        meta_author = soup.find("meta", attrs={"name": "author"})
        if meta_author:
            author = meta_author.get("content", "") or ""

        return [Document(
            page_content=text,
            metadata={"source": file_path, "title": title, "author": author},
        )]


def _load_library_md(file_path: str) -> list[Document]:
    """
    Carrega arquivo .md da biblioteca (AKASHA/Hermes) extraindo campos do frontmatter YAML.
    Separa o bloco YAML do corpo — page_content recebe apenas o texto, sem os metadados.
    Campos extraídos: title, author, date, language, doc_type (= campo `type` do frontmatter).
    """
    try:
        raw_text = open(file_path, encoding="utf-8", errors="ignore").read()
    except OSError as exc:
        raise DocumentLoadError(file_path, str(exc)) from exc

    fm: dict = {}
    body = raw_text
    try:
        import frontmatter as _fm
        post = _fm.loads(raw_text)
        fm   = dict(post.metadata)
        body = post.content
    except Exception:
        pass  # sem frontmatter válido — usa texto bruto

    title          = str(fm.get("title", ""))    or os.path.splitext(os.path.basename(file_path))[0]
    author         = str(fm.get("author", ""))
    date           = str(fm.get("date", ""))
    language       = str(fm.get("language", ""))
    doc_type       = str(fm.get("type", ""))
    fm_source      = str(fm.get("source", ""))   # "mnemosyne_studio" para outputs do Studio

    return [Document(
        page_content=body.strip() or raw_text,
        metadata={
            "source":        file_path,
            "title":         title,
            "author":        author,
            "date":          date,
            "language":      language,
            "doc_type":      doc_type,
            "fm_source":     fm_source,
        },
    )]


def _load_epub(file_path: str) -> list[Document]:
    """
    Carrega um EPUB e retorna 1 Document por capítulo com metadata title/author/chapter.
    Capítulos com menos de 100 chars (capa, índice) são ignorados.

    Raises:
        DocumentLoadError: se ebooklib ou BeautifulSoup não estiver disponível ou falhar.
    """
    try:
        import ebooklib
        from ebooklib import epub
        from bs4 import BeautifulSoup
    except ImportError as exc:
        raise DocumentLoadError(
            file_path,
            f"Dependências em falta para EPUB: {exc}. "
            "Instale com: pip install ebooklib beautifulsoup4",
        ) from exc

    try:
        book = epub.read_epub(file_path, options={"ignore_ncx": True})
    except Exception as exc:
        raise DocumentLoadError(file_path, f"Falha ao abrir EPUB: {exc}") from exc

    title = ""
    author = ""
    try:
        title_meta = book.get_metadata("DC", "title")
        if title_meta:
            title = title_meta[0][0]
        creator_meta = book.get_metadata("DC", "creator")
        if creator_meta:
            author = creator_meta[0][0]
    except Exception:
        pass

    documents: list[Document] = []
    chapter_num = 0

    for item in book.get_items_of_type(ebooklib.ITEM_DOCUMENT):
        try:
            soup = BeautifulSoup(item.get_content(), "lxml")
            text = soup.get_text(separator="\n", strip=True)
        except Exception:
            continue

        if len(text) < 100:
            continue

        chapter_num += 1
        chapter_name = ""
        heading = soup.find(["h1", "h2", "h3"])
        if heading:
            chapter_name = heading.get_text(strip=True)
        if not chapter_name:
            chapter_name = f"Capítulo {chapter_num}"

        documents.append(
            Document(
                page_content=text,
                metadata={
                    "source": file_path,
                    "title": title,
                    "author": author,
                    "chapter": chapter_name,
                    "chapter_num": chapter_num,
                },
            )
        )

    return documents
